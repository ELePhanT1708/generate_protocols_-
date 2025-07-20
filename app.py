from fastapi import FastAPI, UploadFile, File, HTTPException, Response, Form
from fastapi.responses import FileResponse
from docx import Document
import traceback
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
from typing import Dict, List
import re
import os
from collections import defaultdict
import tempfile
import zipfile
import io

from mapping import mapping, mapping_protocol_name
from parse_name import extract_app_info
from replacing_substring import replace_text_with_formatting
import logging
from logging.handlers import RotatingFileHandler
from pathlib import Path


# Настройка логирования
def setup_logging():
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / "app.log"

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        handlers=[
            RotatingFileHandler(log_file, maxBytes=5 * 1024 * 1024, backupCount=3),  # 5 MB на файл, храним 3 версии
            logging.StreamHandler()  # Вывод в консоль
        ]
    )


setup_logging()
logger = logging.getLogger(__name__)

app = FastAPI()

# Путь к исходной заявке
SOURCE_DOC = "636. ООО КХ г. Дятьково.docx"
OOO_TITLE = "Общество с ограниченной ответственностью «Коммунальное хозяйство г. Дятьково»"

# Шаблоны для 5 программ
TEMPLATES = {
    '1': "templates/one_row/00. ПП Шаблон.docx",
    '2': "templates/one_row/00. СИЗ ШАБЛОН.docx",
    '3': "templates/one_row/00. А ШАБЛОН.docx",
    '4': "templates/one_row/00.Б ШАБЛОН.docx",
    '5': "templates/one_row/00. В ШАБЛОН.docx",
}

TEMPLATE_LIST_ATTENDANCE = "templates/one_row/00. УП пустой.docx"


def check_tables_in_file(doc):
    for idx, table in enumerate(doc.tables):
        print(f"\n=== Таблица {idx} ===")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"  Строка {row_idx}: {row_data}")


# 1. Извлечь данные из исходного файла
def parse_applications(path):
    doc = Document(path)
    rows = []
    for table in doc.tables:
        for row in table.rows[1:]:  # Пропускаем заголовок
            cells = row.cells
            try:
                fio = cells[1].text.strip()
                snils = cells[2].text.strip()
                role = cells[3].text.strip()
                programs = re.split(r'[,\s]+', cells[4].text.strip())
                if fio:  # Пропускаем пустые строки
                    rows.append({
                        'fio': fio,
                        'snils': snils,
                        'role': role,
                        'programs': programs
                    })
            except IndexError:
                continue
    return rows


# 2. Сгруппировать сотрудников по номерам программ
def group_by_program(rows):
    program_dict = defaultdict(list)
    for row in rows:
        for program in row['programs']:
            program_dict[program].append(row)
    return program_dict


def clone_row(table, row_idx: int, i: int):
    """
    Клонирует строку с индексом row_idx и возвращает ссылку на ячейки новой строки.
    """
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    if i == 1:
        tbl.remove(tr)
    tbl.append(new_tr)
    return table.rows[-1].cells


def fill_cell(cell, value):
    """
    Чистая вставка текста без лишних переносов, с сохранением форматирования
    """
    cell.text = ""  # Очищаем
    p = cell.paragraphs[0]
    p.clear()


    run = p.add_run(str(value).strip())

    # Устанавливаем шрифт Times New Roman 10
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


@app.post("/generate_protocols/")
async def generate_protocols(
        file: UploadFile = File(...),
        organization_name: str = Form(...)):
    # Логируем начало обработки файла
    logger.info(f"Начало обработки файла: {file.filename}")
    try:
        # Создаем полный путь для сохранения
        save_path = file.filename

        # Сохраняем файл
        with open(save_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)

        # Парсим данные
        logger.info("Парсинг данных из файла...")
        rows = parse_applications(save_path)
        logger.info(f"Найдено {len(rows)} записей")
        grouped_data = group_by_program(rows)
        logger.info(f"Данные сгруппированы по программам: {list(grouped_data.keys())}")
        app_number, org_name = extract_app_info(save_path)
        logger.info(f"Номер заявки: {app_number}, Организация: {org_name}")
        output_dir = tempfile.mkdtemp()
        logger.info(f"Выходная директория: {output_dir}")
        generated_files = []
        for program, people in grouped_data.items():
            logger.info(f"Обработка программы {program} ({len(people)} человек)")
            try:

                doc = Document(TEMPLATES.get(program))

                if int(program) > 5:
                    doc = Document(TEMPLATES['5'])
                    OLD_THEME = (
                        "«Обучение по безопасным методам и приемам выполнения работ повышенной опасности, к которым предъявляются "
                        "дополнительные требования в соответствии с нормативными правовыми актами, содержащими государственные "
                        "нормативные требования охраны труда»"
                    )
                    NEW_THEME = mapping[program]
                    old_protocol_substring = mapping_protocol_name.get('5')
                    new_line = old_protocol_substring.replace("_____", app_number)
                    replace_text_with_formatting(doc, old_protocol_substring, new_line, highlight_substring=new_line)
                    replace_text_with_formatting(doc, OLD_THEME, NEW_THEME)
                else:
                    old_protocol_substring = mapping_protocol_name.get(program)
                    new_line = old_protocol_substring.replace("_____", app_number)
                    replace_text_with_formatting(doc, old_protocol_substring, new_line, highlight_substring=new_line)
                table = doc.tables[1]  # Первая таблица — целевая

                template_row_idx = 1  # ← индекс строки-образца (0 — шапка, 1 — строка для клонирования)

                for i, person in enumerate(people, start=1):
                    cells = clone_row(table, template_row_idx, i)
                    values = [
                        str(f"{i}."),
                        person['fio'],
                        person['role'],
                        organization_name,
                        "удовлетворительно",
                        "Согласно Приложению № 1 к настоящему протоколу",
                        "",  # Дата
                        "",  # Подпись
                    ]

                    for cell, value in zip(cells, values):
                        fill_cell(cell, value)

                # (необязательно) удалить строку-образец:
                # table._tbl.remove(table.rows[template_row_idx]._tr)
                safe_org = re.sub(r'[^\w\s-]', '', org_name).strip().replace(' ', '_')
                output_path = os.path.join(output_dir, f"Протокол_{app_number}_{safe_org}_Программа_{program}.docx")
                doc.save(output_path)
                generated_files.append(output_path)
                logger.info(f" ✅  Файл сохранен: {output_path}")

                ## list attendance
                list_attendance_template = Document(TEMPLATE_LIST_ATTENDANCE)
                old_group_name = "Группа ______________________"
                new_line = f"Группа {app_number}_{safe_org}_Программа_{program}"
                replace_text_with_formatting(list_attendance_template, old_group_name, new_line, highlight_substring=new_line)

                list_attendance_table = list_attendance_template.tables[0]  # Первая таблица — целевая
                template_row_idx_attendance = 3
                for i, person in enumerate(people, start=1):
                    cells = clone_row(list_attendance_table, template_row_idx_attendance, i)
                    values = [
                        str(f"{i}."),
                        person['fio'],
                    ]

                    for cell, value in zip(cells, values):
                        fill_cell(cell, value)

                # (необязательно) удалить строку-образец:
                # table._tbl.remove(table.rows[template_row_idx]._tr)
                safe_org = re.sub(r'[^\w\s-]', '', org_name).strip().replace(' ', '_')
                output_path = os.path.join(output_dir, f"Лист_Посещении_{app_number}_{safe_org}_Программа_{program}.docx")
                list_attendance_template.save(output_path)
                generated_files.append(output_path)
                logger.info(f" ✅  Файл сохранен: {output_path}")
            except Exception as e:
                logger.error(f"Ошибка при обработке программы {program}: {str(e)}\n{traceback.format_exc()}")
        # Создаем ZIP-архив в памяти
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for file_path in generated_files:
                file_name = os.path.basename(file_path)
                zip_file.write(file_path, file_name)

        # Возвращаем ZIP-архив как ответ
        zip_buffer.seek(0)
        return Response(
            content=zip_buffer.getvalue(),
            media_type="application/zip",
            headers={
                "Content-Disposition": "attachment; filename=protocols.zip",
                "Content-Type": "application/zip"
            }
        )
        if not generated_files:
            logger.error("Не удалось сгенерировать ни одного файла")
            raise HTTPException(status_code=400, detail="Не удалось сгенерировать файлы")
    except Exception as e:
        logger.critical(f"Критическая ошибка: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail="Внутренняя ошибка сервера")

    finally:
        if save_path and os.path.exists(save_path):
            try:
                os.unlink(tmp_path)
                logger.info(f"Временный файл удален: {tmp_path}")
            except Exception as e:
                logger.error(f"Ошибка при удалении временного файла: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
