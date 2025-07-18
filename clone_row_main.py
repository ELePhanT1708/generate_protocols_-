
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

import re
import os
from collections import defaultdict

from mapping import mapping, mapping_protocol_name
from parse_name import extract_app_info
from replacing_substring import replace_text_with_formatting

# Путь к исходной заявке
SOURCE_DOC = "636. ООО КХ г. Дятьково.docx"
OOO_TITLE = "Общество с ограниченной ответственностью «Коммунальное хозяйство г. Дятьково»"


# Шаблоны для 5 программ
TEMPLATES = {
    '1': "templates/one_row/00. ПП Шаблон.docx",
    '2': "templates/one_row/00. СИЗ ШАБЛОН.docx",
    '3': "templates/one_row/00. А ШАБЛОН.docx",
    '4': "templates/one_row/00.Б ШАБЛОН.docx",
    '5': "templates/one_row/00. В ШАБЛОН.docx",
}


def check_tables_in_file(doc):
    for idx, table in enumerate(doc.tables):
        print(f"\n=== Таблица {idx} ===")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"  Строка {row_idx}: {row_data}")


# 1. Извлечь данные из исходного файла
def parse_applications(path):
    doc = Document(path)
    rows = []
    for table in doc.tables:
        for row in table.rows[1:]:  # Пропускаем заголовок
            cells = row.cells
            try:
                fio = cells[1].text.strip()
                snils = cells[2].text.strip()
                role = cells[3].text.strip()
                programs = re.split(r'[,\s]+', cells[4].text.strip())
                if fio:  # Пропускаем пустые строки
                    rows.append({
                        'fio': fio,
                        'snils': snils,
                        'role': role,
                        'programs': programs
                    })
            except IndexError:
                continue
    return rows


# 2. Сгруппировать сотрудников по номерам программ
def group_by_program(rows):
    program_dict = defaultdict(list)
    for row in rows:
        for program in row['programs']:
            program_dict[program].append(row)
    return program_dict


def clone_row(table, row_idx: int, i: int):
    """
    Клонирует строку с индексом row_idx и возвращает ссылку на ячейки новой строки.
    """
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    if i == 1:
        tbl.remove(tr)
    tbl.append(new_tr)
    return table.rows[-1].cells


def fill_cell(cell, value):
    """
    Чистая вставка текста без лишних переносов, с сохранением форматирования
    """
    cell.text = ""  # Очищаем
    p = cell.paragraphs[0]
    p.clear()

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(value).strip())

    # Устанавливаем шрифт Times New Roman 10
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def generate_protocols(grouped_data):
    app_number, org_name = extract_app_info(SOURCE_DOC)
    for program, people in grouped_data.items():

        doc = Document(TEMPLATES.get(program))

        if int(program) > 5:
            doc = Document(TEMPLATES['5'])
            OLD_THEME= (
                "«Обучение по безопасным методам и приемам выполнения работ повышенной опасности, к которым предъявляются "
                "дополнительные требования в соответствии с нормативными правовыми актами, содержащими государственные "
                "нормативные требования охраны труда»"
            )
            NEW_THEME = mapping[program]
            old_protocol_substring = mapping_protocol_name.get('5')
            new_line = old_protocol_substring.replace("_____", app_number)
            replace_text_with_formatting(doc, old_protocol_substring, new_line, highlight_substring=new_line)
            replace_text_with_formatting(doc, OLD_THEME, NEW_THEME)
        else:
            old_protocol_substring = mapping_protocol_name.get(program)
            new_line = old_protocol_substring.replace("_____", app_number)
            replace_text_with_formatting(doc, old_protocol_substring, new_line, highlight_substring=new_line)
        table = doc.tables[1]  # Первая таблица — целевая

        template_row_idx = 1  # ← индекс строки-образца (0 — шапка, 1 — строка для клонирования)

        for i, person in enumerate(people, start=1):
            cells = clone_row(table, template_row_idx, i)
            values = [
                str(f"{i}."),
                person['fio'],
                person['role'],
                OOO_TITLE,
                "удовлетворительно",
                "Согласно Приложению № 1 к настоящему протоколу",
                "",  # Дата
                "",  # Подпись
            ]

            for cell, value in zip(cells, values):
                fill_cell(cell, value)

        # (необязательно) удалить строку-образец:
        # table._tbl.remove(table.rows[template_row_idx]._tr)
        safe_org = re.sub(r'[^\w\s-]', '', org_name).strip().replace(' ', '_')
        output_path = f"Протокол_{app_number}_{safe_org}_Программа_{program}.docx"
        doc.save(output_path)
        print(f"✅ Сохранено: {output_path}")


if __name__ == '__main__':
    rows = parse_applications(SOURCE_DOC)
    print(rows)
    grouped = group_by_program(rows)
    print(grouped)
    generate_protocols(grouped)
    # check_tables_in_file(Document(RESULT_FILE))
