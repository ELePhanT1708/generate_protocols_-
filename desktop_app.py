import logging
from logging.handlers import RotatingFileHandler
from pathlib import Path
from tkinter import Tk, filedialog, simpledialog, messagebox
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from copy import deepcopy
import re

import sys
import os
from collections import defaultdict
import tempfile
import zipfile
import traceback

from mapping import mapping, mapping_protocol_name
from parse_name import extract_app_info
from replacing_substring import replace_text_with_formatting


# Настройка логирования
def setup_logging():
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / "desktop_app.log"

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        handlers=[
            RotatingFileHandler(log_file, maxBytes=5 * 1024 * 1024, backupCount=3),  # 5 MB на файл, храним 3 версии
            logging.StreamHandler()  # Вывод в консоль
        ]
    )


def resource_path(relative_path):
    """Возвращает абсолютный путь к ресурсу, работает как при запуске из .py, так и из .exe"""
    try:
        base_path = sys._MEIPASS  # Временная папка PyInstaller
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


setup_logging()
logger = logging.getLogger(__name__)


# Шаблоны для 5 программ
TEMPLATES = {
    '1': resource_path("templates/one_row/00. ПП Шаблон.docx"),
    '2': resource_path("templates/one_row/00. СИЗ ШАБЛОН.docx"),
    '3': resource_path("templates/one_row/00. А ШАБЛОН.docx"),
    '4': resource_path("templates/one_row/00.Б ШАБЛОН.docx"),
    '5': resource_path("templates/one_row/00. В ШАБЛОН.docx"),
}

TEMPLATE_LIST_ATTENDANCE = resource_path("templates/one_row/00. УП пустой.docx")
TEMPLATE_AGREEMENT = resource_path("templates/one_row/00. Шаблон.docx")


def check_tables_in_file(doc):
    for idx, table in enumerate(doc.tables):
        print(f"\n=== Таблица {idx} ===")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"  Строка {row_idx}: {row_data}")


# 1. Извлечь данные из исходного файла
def parse_applications(path):
    doc = Document(path)
    rows = []
    for table in doc.tables:
        for row in table.rows[1:]:  # Пропускаем заголовок
            cells = row.cells
            try:
                fio = cells[1].text.strip()
                snils = cells[2].text.strip()
                role = cells[3].text.strip()
                programs = re.split(r'[,\s]+', cells[4].text.strip())
                if fio:  # Пропускаем пустые строки
                    rows.append({
                        'fio': fio,
                        'snils': snils,
                        'role': role,
                        'programs': programs
                    })
            except IndexError:
                continue
    return rows


# 2. Сгруппировать сотрудников по номерам программ
def group_by_program(rows):
    program_dict = defaultdict(list)
    added_to_program_5 = set()

    for row in rows:
        for program in row['programs']:
            try:
                prog_num = int(program)
            except ValueError:
                continue  # Пропускаем нечисловые значения

            if prog_num > 5:
                # Уникальный идентификатор сотрудника
                identifier = f"{row['fio']}|{row['snils']}"
                if identifier not in added_to_program_5:
                    program_dict['5'].append(row)
                    added_to_program_5.add(identifier)
            else:
                program_dict[str(prog_num)].append(row)

    return program_dict


def clone_row(table, row_idx: int, i: int):
    """
    Клонирует строку с индексом row_idx и возвращает ссылку на ячейки новой строки.
    """
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    if i == 1:
        tbl.remove(tr)
    tbl.append(new_tr)
    return table.rows[-1].cells


def fill_cell(cell, value, alignment='left'):
    """
    Чистая вставка текста без лишних переносов, с сохранением форматирования
    """
    cell.text = ""  # Очищаем
    p = cell.paragraphs[0]
    p.clear()

    run = p.add_run(str(value).strip())

    # Устанавливаем шрифт Times New Roman 10
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if alignment == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if alignment == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if alignment == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Импортируй здесь свои функции


def generate_protocols_from_file(file_path: object, organization_name: object, dogovor: object) -> object:
    try:
        logger.info(f"Начало обработки файла: {file_path}")
        save_path = file_path
        # Получаем директорию, где лежит исходный Excel
        base_dir = os.path.dirname(file_path)

        # Создаем временную папку для генерации DOCX файлов
        output_dir = tempfile.mkdtemp()



        # Парсим данные
        rows = parse_applications(save_path)
        logger.info(f"Найдено {len(rows)} записей")
        grouped_data = group_by_program(rows)
        logger.info(f"Данные сгруппированы по программам: {list(grouped_data.keys())}")
        logger.info(f"Данные сгруппированы по программам: {grouped_data}")
        logger.info(f"Номер заявки: {dogovor}, Организация: {organization_name}")
        logger.info(f"Выходная директория: {output_dir}")
        generated_files = []

        for program, people in grouped_data.items():
            logger.info(f"Обработка программы {program} ({len(people)} человек)")
            try:
                doc = Document(TEMPLATES.get(program))

                if int(program) > 5:
                    doc = Document(TEMPLATES['5'])

                    old_protocol_substring = mapping_protocol_name.get('5')
                    new_line = old_protocol_substring.replace("_____", dogovor)
                    replace_text_with_formatting(doc, old_protocol_substring, new_line, highlight_substring=new_line)
                else:
                    old_protocol_substring = mapping_protocol_name.get(program)
                    new_line = old_protocol_substring.replace("_____", dogovor)
                    replace_text_with_formatting(doc, old_protocol_substring, new_line, highlight_substring=new_line)

                table = doc.tables[1]
                template_row_idx = 1

                for i, person in enumerate(people, start=1):
                    cells = clone_row(table, template_row_idx, i)
                    values = [
                        str(f"{i}."), person['fio'], person['role'],
                        organization_name, "удовлетворительно",
                        "Согласно Приложению № 1 к настоящему протоколу", "", ""
                    ]
                    for cell, value in zip(cells, values):
                        fill_cell(cell, value, 'right')

                safe_org = re.sub(r'[^\w\s-]', '', organization_name).strip().replace(' ', '_')
                output_path = os.path.join(output_dir, f"Протокол_{dogovor}_{safe_org}_Программа_{program}.docx")
                doc.save(output_path)
                generated_files.append(output_path)
                logger.info(f"✅ Файл сохранен: {output_path}")
            except Exception as e:
                logger.error(f"Ошибка при обработке программы {program}: {str(e)}\n{traceback.format_exc()}")

        # ✅ Собираем уникальных сотрудников (по fio + snils)
        unique_attendance = {}
        for row in rows:
            identifier = f"{row['fio']}|{row['snils']}"
            if identifier not in unique_attendance:
                unique_attendance[identifier] = row

        # ✅ Готовим Лист посещения
        logger.info(f'Уникальных сотрудников :{unique_attendance}')
        try:
            list_attendance_template = Document(TEMPLATE_LIST_ATTENDANCE)
            new_line = f"Группа {dogovor}_{safe_org}"
            replace_text_with_formatting(
                list_attendance_template,
                "Группа ______________________",
                new_line,
                highlight_substring=new_line
            )

            list_attendance_table = list_attendance_template.tables[0]
            template_row_idx_attendance = 3

            for i, person in enumerate(unique_attendance.values(), start=1):
                cells = clone_row(list_attendance_table, template_row_idx_attendance, i)
                values = [str(f"{i}."), person['fio']]
                for cell, value in zip(cells, values):
                    fill_cell(cell, value, 'left')

            output_path = os.path.join(output_dir, f"Лист_Посещения_{dogovor}_{safe_org}.docx")
            list_attendance_template.save(output_path)
            generated_files.append(output_path)
            logger.info(f"✅ Общий Лист посещения сохранен: {output_path}")

        except Exception as e:
            logger.error(f"Ошибка при создании общего Листа посещения: {str(e)}\n{traceback.format_exc()}")

        # ✅ Готовим Согласия
        logger.info(f'Уникальных сотрудников :{unique_attendance}')
        try:
            list_agreement_template = Document(TEMPLATE_AGREEMENT)
            new_line = f"_________{safe_org}_________"
            replace_text_with_formatting(
                list_agreement_template,
                "__________________________________________________________________,",
                new_line,
                highlight_substring=new_line
            )

            list_agreement_table = list_agreement_template.tables[0]
            template_row_idx_attendance = 1

            for i, person in enumerate(unique_attendance.values(), start=1):
                cells = clone_row(list_agreement_table, template_row_idx_attendance, i)
                values = [str(f"{i}."), person['fio']]
                for cell, value in zip(cells, values):
                    fill_cell(cell, value, 'left')

            output_path = os.path.join(output_dir, f"Согласие_{dogovor}_{safe_org}.docx")
            list_agreement_template.save(output_path)
            generated_files.append(output_path)
            logger.info(f"✅ Общее согласие сохранено: {output_path}")

        except Exception as e:
            logger.error(f"Ошибка при создании общего Листа согласии: {str(e)}\n{traceback.format_exc()}")


        # Архивируем
        # Указываем финальный путь архива
        zip_filename = f"{dogovor}.zip"
        zip_path = os.path.join(base_dir, zip_filename)
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file_path in generated_files:
                zipf.write(file_path, os.path.basename(file_path))

        logger.info(f"Архив создан: {zip_path}")
        messagebox.showinfo("Успех", f"Протоколы сгенерированы:\n{zip_path}")

    except Exception as e:
        logger.critical(f"Критическая ошибка: {str(e)}\n{traceback.format_exc()}")
        messagebox.showerror("Ошибка", f"Ошибка генерации: {str(e)}")


def main_gui():
    root = Tk()
    root.withdraw()  # Не показывать основное окно

    file_path = filedialog.askopenfilename(
        title="Выберите файл для загрузки",
        filetypes=[("Word", "*.docx")]
    )

    if not file_path:
        messagebox.showwarning("Внимание", "Файл не выбран.")
        return

    organization_name = simpledialog.askstring("Организация", "Введите название организации:")
    if not organization_name:
        messagebox.showwarning("Внимание", "Название организации не указано.")
        return

    dogovor = simpledialog.askstring("Номер Договора", "Введите номер договора:")
    if not dogovor:
        messagebox.showwarning("Внимание", "Номер Договора не указан.")
        return

    generate_protocols_from_file(file_path, organization_name, dogovor)


if __name__ == "__main__":
    main_gui()
