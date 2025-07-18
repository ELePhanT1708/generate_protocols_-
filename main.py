from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import re
import os
from collections import defaultdict

# Путь к исходной заявке
SOURCE_DOC = "636. ООО КХ г. Дятьково.docx"
TEMPLATE_1 = "templates/00. ПП Шаблон.docx"
TEMPLATE_2 = "templates/00. СИЗ ШАБЛОН.docx"
TEMPLATE_3 = "templates/00. А ШАБЛОН.docx"
TEMPLATE_4 = "templates/00.Б ШАБЛОН.docx"
TEMPLATE_5 = "templates/00. В ШАБЛОН.docx"

# Шаблоны для 5 программ
TEMPLATES = {
    '1': "templates/00. ПП Шаблон.docx",
    '2': "templates/00. СИЗ ШАБЛОН.docx",
    '3': "templates/00. А ШАБЛОН.docx",
    '4': "templates/00.Б ШАБЛОН.docx",
    '5': "templates/00. В ШАБЛОН.docx",
}


def check_tables_in_file(doc):
    for idx, table in enumerate(doc.tables):
        print(f"\n=== Таблица {idx} ===")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"  Строка {row_idx}: {row_data}")


# 1. Извлечь данные из исходного файла
def parse_applications(path):
    doc = Document(path)
    rows = []
    for table in doc.tables:
        for row in table.rows[1:]:  # Пропускаем заголовок
            cells = row.cells
            try:
                fio = cells[1].text.strip()
                snils = cells[2].text.strip()
                role = cells[3].text.strip()
                programs = re.split(r'[,\s]+', cells[4].text.strip())
                if fio:  # Пропускаем пустые строки
                    rows.append({
                        'fio': fio,
                        'snils': snils,
                        'role': role,
                        'programs': programs
                    })
            except IndexError:
                continue
    return rows


# 2. Сгруппировать сотрудников по номерам программ
def group_by_program(rows):
    program_dict = defaultdict(list)
    for row in rows:
        for program in row['programs']:
            if program in TEMPLATES:
                program_dict[program].append(row)
    return program_dict


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Установка границ для ячейки
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 8)))
            edge_el.set(qn('w:space'), str(edge_data.get('space', 0)))
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)


# Установка внутренних отступов в ячейке (в twips, 0 = без отступов)
def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    cellMar = tcPr.find(qn('w:tcMar'))
    if cellMar is None:
        cellMar = OxmlElement('w:tcMar')
        tcPr.append(cellMar)

    for margin_type, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = cellMar.find(qn(f'w:{margin_type}'))
        if node is None:
            node = OxmlElement(f'w:{margin_type}')
            cellMar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


# Установка вертикального выравнивания вверх
def set_cell_vertical_alignment_top(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'top')
    tcPr.append(vAlign)


# Главная функция генерации протоколов
def generate_protocols(grouped_data):
    for program, people in grouped_data.items():
        doc = Document(TEMPLATES[program])
        table = doc.tables[1]  # предполагаем, что нужная таблица — вторая
        for i, person in enumerate(people, start=1):
            row = table.add_row().cells
            values = [
                str(i),
                person['fio'].replace('\n', ' ').strip(),
                person['snils'].replace('\n', ' ').strip(),
                person['role'].replace('\n', ' ').strip(),
                "удовлетворительно",  # Результат проверки
                "Согласно Приложению № 1 к настоящему протоколу",  # Рег. номер
                "",  # Дата
                "",  # Подпись
            ]

            for cell, value in zip(row, values):
                # Чистая вставка текста
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.style = doc.styles['Normal']
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1

                run = paragraph.add_run(value)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                set_cell_border(cell,
                                top={"val": "single", "sz": 6, "color": "000000"},
                                bottom={"val": "single", "sz": 6, "color": "000000"},
                                left={"val": "single", "sz": 6, "color": "000000"},
                                right={"val": "single", "sz": 6, "color": "000000"},
                                )
                set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
                set_cell_vertical_alignment_top(cell)

        output_path = f"Протокол_Программа_{program}.docx"
        doc.save(output_path)
        print(f"✅ Сохранено: {output_path}")


if __name__ == '__main__':
    # check_tables_in_file(Document(SOURCE_DOC))
    # check_tables_in_file(Document(TEMPLATE_1))
    # check_tables_in_file(Document(TEMPLATE_2))
    # check_tables_in_file(Document(TEMPLATE_3))
    # check_tables_in_file(Document(TEMPLATE_4))
    # check_tables_in_file(Document(TEMPLATE_5))
    rows = parse_applications(SOURCE_DOC)
    print(rows)
    grouped = group_by_program(rows)
    print(grouped)
    generate_protocols(grouped)
