import re
from docx import Document
from docx.shared import Pt

from mapping import mapping_protocol_name


def replace_text_with_formatting(doc, old_text, new_text, highlight_substring=None):
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)

        match = re.search(re.escape(old_text), full_text)
        if not match:
            continue

        start, end = match.span()

        current_pos = 0
        before_runs = []
        match_runs = []
        after_runs = []

        for run in paragraph.runs:
            run_len = len(run.text)
            if current_pos + run_len <= start:
                before_runs.append(run)
            elif current_pos >= end:
                after_runs.append(run)
            else:
                match_runs.append(run)
            current_pos += run_len

        # Сохраняем пробелы до и после
        before_text = full_text[:start]
        after_text = full_text[end:]

        add_space_before = before_text.endswith(' ')
        add_space_after = after_text.startswith(' ')

        # Удаляем все run'ы
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

        # Копирование форматирования
        def append_run(p, ref_run, text, make_bold=False):
            new_run = p.add_run(text)
            new_run.bold = make_bold or ref_run.bold
            new_run.italic = ref_run.italic
            new_run.underline = ref_run.underline
            new_run.font.name = ref_run.font.name
            new_run.font.size = ref_run.font.size
            return new_run

        # Вставляем back
        for run in before_runs:
            append_run(paragraph, run, run.text)

        # Формируем final_new_text с жирным, если надо
        ref_run = match_runs[0] if match_runs else paragraph.add_run()

        if add_space_before:
            append_run(paragraph, ref_run, ' ')

        if highlight_substring and highlight_substring in new_text:
            # Разбиваем текст на части
            before, match, after = new_text.partition(highlight_substring)
            if before:
                append_run(paragraph, ref_run, before)
            append_run(paragraph, ref_run, match, make_bold=True)
            if after:
                append_run(paragraph, ref_run, after)
        else:
            append_run(paragraph, ref_run, new_text)

        if add_space_after:
            append_run(paragraph, ref_run, ' ')

        for run in after_runs:
            append_run(paragraph, run, run.text)

if __name__ == '__main__':
    SOURCE_FILE = "templates/empty/00. В ШАБЛОН.docx"
    OUTPUT_FILE = "output.docx"

    doc = Document(SOURCE_FILE)

    replace_text_with_formatting(doc, mapping_protocol_name['5'], 'ПРОТОКОЛ №В-636', highlight_substring="ПРОТОКОЛ №В-636")

    doc.save(OUTPUT_FILE)
    print("✅ Замена завершена. Файл сохранён:", OUTPUT_FILE)
