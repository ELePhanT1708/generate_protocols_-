from docx import Document
import re
from mapping import mapping, mapping_protocol_name


def replace_text_with_formatting(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)

        match = re.search(re.escape(old_text), full_text)
        if not match:
            continue

        start, end = match.span()

        current_pos = 0
        before_runs = []
        match_runs = []
        after_runs = []

        for run in paragraph.runs:
            run_len = len(run.text)
            if current_pos + run_len <= start:
                before_runs.append(run)
            elif current_pos >= end:
                after_runs.append(run)
            else:
                match_runs.append(run)
            current_pos += run_len

        # Сохраняем пробелы до и после, если они есть
        before_text = full_text[:start]
        after_text = full_text[end:]

        add_space_before = before_text.endswith(' ')
        add_space_after = after_text.startswith(' ')

        # Удаляем все run'ы
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

        # Функция для копирования стиля
        def append_run(p, ref_run, text):
            new_run = p.add_run(text)
            new_run.bold = ref_run.bold
            new_run.italic = ref_run.italic
            new_run.underline = ref_run.underline
            new_run.font.name = ref_run.font.name
            new_run.font.size = ref_run.font.size
            return new_run

        # Вставляем back
        for run in before_runs:
            append_run(paragraph, run, run.text)

        # Формируем финальный заменённый текст
        final_new_text = ""
        if add_space_before:
            final_new_text += " "
        final_new_text += new_text
        if add_space_after:
            final_new_text += " "

        ref_run = match_runs[0] if match_runs else paragraph.add_run()
        append_run(paragraph, ref_run, final_new_text)

        for run in after_runs:
            append_run(paragraph, run, run.text)


if __name__ == '__main__':
    SOURCE_FILE = "templates/empty/00. В ШАБЛОН.docx"
    OUTPUT_FILE = "output.docx"

    doc = Document(SOURCE_FILE)

    replace_text_with_formatting(doc, mapping_protocol_name['5'], 'ПРОТОКОЛ №В-636')

    doc.save(OUTPUT_FILE)
    print("✅ Замена завершена. Файл сохранён:", OUTPUT_FILE)
